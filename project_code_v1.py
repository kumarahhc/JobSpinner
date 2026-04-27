from datetime import datetime
import pandas as pd
import autogen
import os
import fitz  # PyMuPDF
from docx import Document
from typing import Annotated
from duckduckgo_search import DDGS
import json

# 1. Configuration for OpenRouter Proxy
# Replace the base_url with your specific AWS proxy URL
config_list = [
    {
        "model": "openai/gpt-4.1-mini", 
        "base_url": "https://5f5832nb90.execute-api.eu-central-1.amazonaws.com/v1", # Your provided AWS proxy URL
        "api_key": "not-needed",       # Proxy handles auth
    }
]

llm_config = {"config_list": config_list, "cache_seed": None}

# --- TOOLS ---

# def search_jobs(query: Annotated[str, "The job search query including role and location"]) -> str:
#     """Searches for job postings using DuckDuckGo."""
#     with DDGS() as ddgs:
#         results = [r for r in ddgs.text(query, max_results=5)]
#         return json.dumps(results) if results else "No jobs found."

# def save_to_log(job_details: Annotated[str, "JSON string of job details to save"]) -> str:
#     """Saves the job application info to a local JSON file."""
#     try:
#         data = json.loads(job_details)
#         with open("jobs_log.json", "a") as f:
#             f.write(json.dumps(data) + "\n")
#         return "Successfully logged to jobs_log.json"
#     except Exception as e:
#         return f"Error logging: {e}"

def read_text_file(file_path):
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    return "Standard professional cover letter template."


def save_cover_letter_docx(content: str, company_name: str = "Company") -> str:
    """Saves the cover letter text into a formatted .docx file."""
    doc = Document()
    
    # Add a title and date
    doc.add_heading('Cover Letter', 0)
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d"))
    
    # Add the AI generated content
    doc.add_paragraph(content)
    
    # Create a safe filename
    filename = f"Cover_Letter_{company_name.replace(' ', '_')}.docx"
    doc.save(filename)
    
    return f"File saved successfully as {filename}"   

def save_latex_file(latex_code: Annotated[str, "The complete LaTeX source code"], 
                    filename: Annotated[str, "The name of the file (e.g., cover_letter.tex)"]) -> str:
    """Saves the provided LaTeX code into a .tex file."""
    try:
        with open(filename, "w", encoding="utf-8") as f:
            f.write(latex_code)
        return f"Successfully saved LaTeX to {filename}"
    except Exception as e:
        return f"Error saving file: {e}"

def update_job_excel(
    company: Annotated[str, "Company name"],
    role: Annotated[str, "Job title/role"],
    url: Annotated[str, "URL of the job posting"],
    status: Annotated[str, "Status (e.g., Applied, Tailored)"] = "Applied"
) -> str:
    """Saves or updates job application details in a local Excel file."""
    file_path = "job_application_status.xlsx"
    current_date = datetime.now().strftime("%Y-%m-%d")
    
    new_data = {
        "Date": [current_date],
        "Company": [company],
        "Role": [role],
        "URL": [url],
        "Status": [status]
    }
    new_df = pd.DataFrame(new_data)

    if os.path.exists(file_path):
        # Read existing data
        existing_df = pd.read_excel(file_path)
        
        # Check if the job URL already exists to prevent duplicates
        if url in existing_df['URL'].values:
            return f"Job at {company} already exists in the tracker. No duplicate added."
        
        # Append and save
        updated_df = pd.concat([existing_df, new_df], ignore_index=True)
        updated_df.to_excel(file_path, index=False)
        return f"Appended new job at {company} to {file_path}."
    else:
        # Create new file
        new_df.to_excel(file_path, index=False)
        return f"Created new file and saved job at {company}."

# --- AGENT DEFINITIONS ---


# Updated User Proxy with safe termination check
user_proxy = autogen.UserProxyAgent(
    name="User_Proxy",
    human_input_mode="ALWAYS", 
    max_consecutive_auto_reply=1,
    # This check now safely handles NoneType content
    is_termination_msg=lambda x: x.get("content") is not None and x.get("content").rstrip().endswith("TERMINATE"),
    code_execution_config=False, 
)
# # Agent 1: The Scout (Researcher)
# scout = autogen.AssistantAgent(
#     name="Scout",
#     system_message="You are a job researcher. Find jobs matching the user's criteria. "
#                    "Use the search_jobs tool. Return the details of the best 3 matches. "
#                    "After providing data, say TERMINATE.",
#     llm_config=llm_config,
# )

# Agent 2: The Tailor (Writer)
tailor = autogen.AssistantAgent(
    name="Tailor",
    system_message="""You are a career coach. 
    1. Analyze the CV and Job Description provided and provide strengths and weaknesses for the applicant in relation to the job in bullet points. Also provide a value for matching precentage of the users skills and job requirements. Your can present this in visually appealing way using emojis and formatting.
    Ask for the user's input to proceed with writing the cover letter.
    2. Write a high-quality cover letter empahsising the applicant's qualifications and fit for the role wihout haulicinating. LIMIT the letter to ONE A4 PAGE. Present it user and Ask for the user's approval before finalizing and writing the letter to a file.
    3. IMPORTANT: Based on user input file type, use the 'save_cover_letter_docx' tool to save the final letter to a docx file or use the 'save_latex_file' tool to save it as a .tex file. 
    When usig .tex, Ensure special characters like '&' or '%' are escaped (e.g., \&). Do Not provide the tex file output on screen since it is not readble, instead directly save it .tex file without showing the content on screen. 
    Provide the company name to the tool for the filename.
    4. Say TERMINATE once the file is saved. And handover to the Registrar to log the application details.""",
    llm_config=llm_config,
)

# Agent definition
registrar = autogen.AssistantAgent(
    name="Registrar",
    system_message="""You are a Data Manager. 
    1. Extract the Company, Role, and URL from the conversation. 
    2. Use the 'update_job_excel' tool to save the information.
    3. Say TERMINATE once the data is saved.""",
    llm_config=llm_config,
)


# --- REGISTER TOOLS ---
# autogen.agentchat.register_function(
#     search_jobs,
#     caller=scout,
#     executor=user_proxy,
#     description="Tool to search for job vacancies on the web",
# )


autogen.agentchat.register_function(
    save_cover_letter_docx,
    caller=tailor,        # The Tailor decides to call it
    executor=user_proxy,  # The User_Proxy executes it on the local machine
    description="Saves the generated cover letter to a .docx file on disk.",
)
autogen.agentchat.register_function(
    save_latex_file,
    caller=tailor,
    executor=user_proxy,
    description="Saves the generated LaTeX cover letter to a local file",
)

autogen.agentchat.register_function(
    update_job_excel,
    caller=registrar,
    executor=user_proxy,
    description="Logs job application details to an Excel file with duplicate prevention",
)


# --- ORCHESTRATION: SEQUENTIAL CHAT ---

def run_job_assistant(job_search_criteria, cv_text):
    # 1. Scout finds the jobs
    # res_scout = user_proxy.initiate_chat(
    #     scout, 
    #     message=f"Search for: {job_search_criteria}",
    #     summary_method="last_msg"
    # )
    
    # 2. Tailor uses the CV text to write the letter
    # We combine the CV text and the job findings into one message
    tailor_instruction = f"""
    Here is my CV content:
    {cv_text}
    
    Based on this CV, draft a cover letter according to the job requirements and details
    for the following job found:
    {job_search_criteria}
    """
    
    user_proxy.initiate_chat(
        tailor, 
        message=tailor_instruction,
        summary_method="last_msg"
    )
    clean_job_info = job_search_criteria.replace("TERMINATE", "").strip()
    # Step 3: Registrar Logs the Data
    user_proxy.initiate_chat(
        registrar, 
        message=f"Log this application detail: {clean_job_info}",
    )


def extract_text_from_file(file_path):
    """Extracts text from PDF or DOCX files."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text
    
    elif ext == ".docx":
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    else:
        return "Unsupported file format."



if __name__ == "__main__":
    # Define the path to your local text file
    file_path = "job_requirements.txt" 
    cv_path = "my_cv.pdf"  # or "my_cv.docx"
    cl_path = "cover_letter_format.txt"
    
    if os.path.exists(cv_path):
        cv_content = extract_text_from_file(cv_path)

    # Check if the file exists to avoid errors
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            # Read the content and strip whitespace
            user_goal_from_file = f.read().strip()
            
        print(f"--- Reading requirements from {file_path} ---")
        
        # Pass the file content into your main function
        run_job_assistant(user_goal_from_file, cv_content)
    else:
        print(f"Error: {file_path} not found in the current folder.")
        # Fallback to manual input if file is missing
        # run_job_assistant("Python Developer in Helsinki")
    
