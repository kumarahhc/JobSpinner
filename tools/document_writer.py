import os
from typing import Annotated
from datetime import datetime
import pandas as pd
from docx import Document
from config import OUTPUT_FOLDER

def update_job_excel(
    company: Annotated[str, "Company name"],
    role: Annotated[str, "Job title/role"],
    url: Annotated[str, "URL of the job posting"],
    status: Annotated[str, "Status (e.g., Applied, Tailored)"] = "Applied"
) -> str:
    """
    Saves or updates applied job application details in a local Excel file.
    Args:
        company (str): Company name.
        role (str): Job title/role.
        url (str): URL of the job posting.
        status (str, optional): Status of the application. Defaults to "Applied".
        Returns:
            str: A message indicating the result of the operation.
    """
    file_name = "job_application_status.xlsx"
    current_date = datetime.now().strftime("%Y-%m-%d")
    
    new_data = {
        "Date": [current_date],
        "Company": [company],
        "Role": [role],
        "URL": [url],
        "Status": [status]
    }
    new_df = pd.DataFrame(new_data)
    file_path = os.path.join(OUTPUT_FOLDER, file_name)
    if os.path.exists(file_path):
        # Read existing data
        existing_df = pd.read_excel(file_path)
        
        # Check if the job URL already exists to prevent duplicates
        if url in existing_df['URL'].values:
            return f"Job at {company} already exists in the tracker. No duplicate added."
        
        # Append and save
        updated_df = pd.concat([existing_df, new_df], ignore_index=True)
        updated_df.to_excel(file_path, index=False)
        return f"Appended new job at {company} to {file_path}."
    else:
        # Create new file
        new_df.to_excel(file_path, index=False)
        return f"Created new file and saved job at {company}."
    

def save_cover_letter_docx(content: str, company_name: str = "Company") -> str:
    """
    Saves the cover letter text into a formatted .docx file.
    Args:
        content (str): The text content of the cover letter.
        company_name (str, optional): The name of the company for filename. Defaults to "Company".
    Returns:
        str: A message indicating the result of the operation.
    """
    doc = Document()
    
    # Add a title and date
    doc.add_heading('Cover Letter', 0)
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d"))
    
    # Add the AI generated content
    doc.add_paragraph(content)
    
    # Create a safe filename
    filename = f"Cover_Letter_{company_name.replace(' ', '_')}.docx"
    file_path = os.path.join(OUTPUT_FOLDER, filename)
    doc.save(file_path)
    
    return f"File saved successfully as {filename}"   

def save_latex_file(latex_code: Annotated[str, "The complete LaTeX source code"], 
                    filename: Annotated[str, "The name of the file (e.g., cover_letter.tex)"]) -> str:
    """
    Saves the provided LaTeX code into a .tex file.
    Args:
        latex_code (str): The complete LaTeX source code.
        filename (str): The name of the file (e.g., cover_letter.tex).
    Returns:
        str: A message indicating the result of the operation.
    """
    file_path = os.path.join(OUTPUT_FOLDER, filename)
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(latex_code)
        return f"Successfully saved LaTeX to {file_path}"
    except Exception as e:
        return f"Error saving file: {e}"